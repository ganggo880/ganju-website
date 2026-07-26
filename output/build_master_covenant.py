import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_master_covenant():
    doc = docx.Document()
    
    # Page Margins (Normal: 1 inch = 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Normal Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = '標楷體'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Dark Charcoal
    style_normal.paragraph_format.line_spacing = 1.3
    style_normal.paragraph_format.space_after = Pt(6)

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_company = p_title.add_run("港居不動產開發有限公司\n")
    run_company.font.name = '微軟正黑體'
    run_company.font.size = Pt(14)
    run_company.font.bold = True
    run_company.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Navy Blue #1A365D

    run_title = p_title.add_run("通用標準住戶公約 (Master 版)")
    run_title.font.name = '微軟正黑體'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    # Sub-header notice
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("【本公約為房屋租賃契約書之重要附件，具同等法律效力】")
    run_sub.font.name = '微軟正黑體'
    run_sub.font.size = Pt(10)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0xC5, 0xA0, 0x59) # Champagne Gold #C5A059
    p_sub.paragraph_format.space_after = Pt(18)

    # 27 Core Clauses
    clauses = [
        ("1. 禁菸酒檳榔與秩序維護", "在室內及公共區域全面禁止吸菸、檳榔及喝酒（如有亂丟菸蒂、檳榔渣、瓶罐、垃圾或喧嘩、鬧事等行為，經查證屬實依規處置）。"),
        ("2. 公共區域禁止堆放私人物品", "樓梯走廊、頂樓、公共通道禁止擺放私人物品，違者視同無主廢棄物逕行清除丟棄，以免妨礙通行與消防安全。"),
        ("3. 嚴禁違法行為與無條件解約條款", "本層樓嚴禁任何違法行為（如色情行業、吸毒、拉K、販毒、賭博等），一經發現立即報警處理。租客在租約存續期間，如有違反公序良俗或任何違法行為（包括但不限於租客本人、同住人或訪客之行為），承租人（乙方）同意無條件立即解約，將房屋恢復原狀返還出租人（甲方），並自願放棄先訴抗辯權。"),
        ("4. 屋內吸菸/拜拜香薰罰則與陽台規範", "屋內及陽台保障大家安全。依菸害防制法第17條及公寓大廈管理條例第23條第1項規定，屋內及後陽台窗嚴禁吸煙、電子煙、加熱菸、點香、香盤、香薰，並禁止由視窗丟棄垃圾或煙頭。違者第一次扣罰新台幣10,000元整（清洗冷氣機、窗簾及屋內除臭費用）；違規第二次則視為重大違約終止契約，由甲方認定事實後三日內搬家並清空屋內，未清空者視同廢棄物處理，乙方不得異議。（乙方簽名：__________________）"),
        ("5. 寵物禁止與罰則", "屋內及公共區域嚴禁飼養與攜帶寵物進入，違者單次扣罰新台幣10,000元整（清洗冷氣機、窗簾及屋內除臭費用），且乙方須立即將寵物帶離租屋處。"),
        ("6. 安寧維護與夜間時段規範", "依據噪音管制法規定，夜間時段（晚上10點至清晨7點）住宅區不得超過平均50分貝。夜間動作、談話音量請放輕，禁止拖行重物。來訪親友應於夜間時段前離開，禁止邀請非承租人回租屋處聚會過夜。"),
        ("7. 水管馬桶嚴禁拋棄異物", "嚴禁將溼式衛生紙（含標榜環保水解材質）、濕紙巾、面紙、貓砂、衛生棉、骨頭菜渣、牙籤、牙線棒及其他異物丟入水管與馬桶，廚餘尤其嚴禁倒入。若因乙方違反規定導致水管或馬桶阻塞，經查屬實乙方須負擔全額疏通及修繕賠償費用，乙方不得異議。"),
        ("8. 浴室防霉與水氣排除", "浴室使用完畢後請開門散去水氣或開啟除濕機/排風扇，以避免牆面及天花板發霉。"),
        ("9. 室內常保通風與壁癌防範", "室內須常保通風良好。若未開窗通風致室內潮濕、牆壁壁癌、設備發霉或電器短路，相關損修費用由乙方負擔。"),
        ("10. 設備故障報修程序", "入住後設備若有損壞故障，請乙方詳細敘述並拍照錄影通知甲方維修，乙方並應配合維修時間。"),
        ("11. 入住前7天管線確認責任歸屬", "租賃期間開始7天內，乙方應先確認馬桶、洗臉盆、流理台水槽及排水口功能是否正常，如有異狀應立即提出由甲方處理。入住7天後若發生堵塞，視為乙方使用不當，由乙方負擔疏通責任。"),
        ("12. 鏡面與玻璃水垢維護", "浴室鏡子及乾濕分離玻璃牆，使用後請隨手擦乾，以免形成無法清除之硬化水垢。"),
        ("13. 冷氣濾網定期清潔", "冷氣機須定期清潔濾網。若因未清洗濾網致冷氣機損壞，修復費用由乙方負擔。"),
        ("14. 洗衣機使用規範與禁用時段", "洗衣機夜間禁用時段為晚上10點至清晨8點，以免運轉噪音影響他人作息。棉絮濾網須定期清潔。"),
        ("15. 社區寧靜與神明廳禁設", "保持社區寧靜，不使音響達於戶外。租屋處內嚴禁設立神明廳、神壇或進行各類焚香祭拜。"),
        ("16. 專人開門車馬費", "遺失或忘帶鑰匙需人員到場開門者，每次酌收車馬費新台幣500元整；晚間6點後派員每次酌收車馬費新台幣1,000元整。"),
        ("17. 磁扣與鑰匙遺失補發費", "大門磁扣若遺失，每支補發300元整；鐵製鑰匙遺失，每支補發300元整（補發需5個工作天）。"),
        ("18. 用途限制", "嚴禁將房屋出租、轉租或作為不正當、違法之行業使用。"),
        ("19. 守望相助", "應發揮守望相助之精神，一發現有意外事故或其他反常事件，應立即聯絡甲方或報警處理。"),
        ("20. 牆面保護與賠償標準", "牆面不破壞不釘鉤物品。如要吊掛物品僅得使用3M無痕掛鉤，不可使用雙面膠、泡綿膠等難以清潔產品。退租前須恢復原狀，如有損壞或殘膠，每處賠償500元整，二處賠償1,000元整，以此類推。"),
        ("21. 緊急事件授權破門", "租賃範圍內發生火災、緊急救護或其他緊急事件時，乙方同意轄區警員、消防員、鄰里長會同甲方得破壞門窗逕行進入管控現場及排除緊急狀況，乙方不得異議。"),
        ("22. 愛惜生命與凶宅違約金條款", "乙方應愛惜生命，不得有自殘、自傷或傷人之行為，如有違者，甲方得提前終止租約。屋內如發生非自然死亡之情事致房屋成為凶宅者，乙方（或其繼承人/連帶保證人）須負擔懲罰性違約金新台幣300萬元整。甲方亦擔保房屋未發生非自然死亡之情事，如有違者，乙方得解除本契約，若因此造成乙方損失，甲方並應賠償。"),
        ("23. 個人衛生與床單規格", "個人衛生考量，請務必鋪設床單。床單、枕頭屬於個人衛生用品由住戶自備，標準雙人床尺寸＝5尺(152公分) x 6.2尺(188公分)。"),
        ("24. 信件包裹處理", "本棟樓無代收信件及包裹服務，請自行至一樓信箱區收取信件。"),
        ("25. 退租清潔標準與費用", "退租時應將房間遷空打掃乾淨交還甲方。床底清潔、電視櫃整體、浴室水垢、冰箱退冰完成為清潔重點。若未遷空清潔恢復原樣，或冰箱未退冰完成，甲方得酌收清潔費，每處新台幣600元整。"),
        ("26. 違約效果與押金沒收", "以上各條約住戶均須嚴格遵守，如有違反者視同違約，甲方有權立即解約並沒收押金，且得隨時檢查房屋現況。"),
        ("27. 公約效力與簽署承諾", "本住戶公約為租賃契約附件之一，乙方簽署即表示充分理解並同意遵守上述所有條款規定。")
    ]

    for title, detail in clauses:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        
        # Title Part (Bold Navy Blue)
        run_t = p.add_run(f"{title}：")
        run_t.font.name = '微軟正黑體'
        run_t.font.size = Pt(11)
        run_t.font.bold = True
        run_t.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        # Detail Part
        run_d = p.add_run(detail)
        run_d.font.name = '標楷體'
        run_d.font.size = Pt(11)
        
        # Highlight strict penalties with Red if applicable
        if "10,000" in detail or "300萬元" in detail or "無條件解約" in detail:
            # We can give key clauses extra emphasis
            pass

    # Signature Block at the end
    p_sig_title = doc.add_paragraph()
    p_sig_title.paragraph_format.space_before = Pt(18)
    p_sig_title.paragraph_format.space_after = Pt(12)
    run_st = p_sig_title.add_run("【立約簽署欄】")
    run_st.font.name = '微軟正黑體'
    run_st.font.size = Pt(12)
    run_st.font.bold = True
    run_st.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.line_spacing = 1.6
    
    r_sig = p_sig.add_run(
        "出租人 (甲方)： 港居不動產開發有限公司\n"
        "代表人 / 管理經理： ______________________ (簽章)\n"
        "聯絡電話： 0968-863-880 / 07-791-2288\n\n"
        "承租人 (乙方)： ______________________ (簽名蓋章)\n"
        "身分證字號： _________________________\n"
        "聯絡電話： ___________________________\n\n"
        "簽署日期： 中華民國 ________ 年 ____ 月 ____ 日"
    )
    r_sig.font.name = '微軟正黑體'
    r_sig.font.size = Pt(11)

    doc.save(r"g:\我的雲端硬碟\ai agent\Matt Pocock skills\output\港居不動產_通用標準住戶公約_Master版.docx")
    print("Master版 Docx 產出成功！")

if __name__ == "__main__":
    create_master_covenant()
